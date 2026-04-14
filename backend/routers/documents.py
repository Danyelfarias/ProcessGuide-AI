import io
import os
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import Optional, List

from database.db import get_db, Document, Process
from services.export_service import export_process_docx

router = APIRouter()


def _extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(p.extract_text() or "" for p in reader.pages)

    if ext in ("docx",):
        import docx
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)

    if ext in ("md", "txt"):
        return content.decode("utf-8", errors="ignore")

    raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    process_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
):
    raw = await file.read()
    text = _extract_text(file.filename, raw)

    doc = Document(
        process_id=process_id,
        filename=file.filename,
        content=text,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    if process_id:
        process = db.query(Process).filter(Process.id == process_id).first()
        if process:
            process.source_doc = file.filename
            db.commit()

    return {"id": doc.id, "filename": doc.filename, "characters": len(text)}


@router.get("/process/{process_id}")
def list_documents(process_id: int, db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.process_id == process_id).all()
    return [{"id": d.id, "filename": d.filename, "uploaded_at": d.uploaded_at} for d in docs]


@router.get("/{document_id}/text")
def get_document_text(document_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"id": doc.id, "filename": doc.filename, "content": doc.content}


@router.delete("/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(doc)
    db.commit()
    return {"message": "deleted", "id": document_id}


@router.get("/export/{process_id}/docx")
def export_docx(process_id: int, db: Session = Depends(get_db)):
    process = db.query(Process).filter(Process.id == process_id).first()
    if not process:
        raise HTTPException(status_code=404, detail="Process not found")

    path = export_process_docx(process.name, process.flow_json)
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{process.name.replace(' ', '_')}.docx",
    )
